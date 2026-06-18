import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ===== 页面设置 =====
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

DARK = RGBColor(0x3A, 0x3F, 0x47)
GRAY = RGBColor(0x88, 0x88, 0x88)
TEXT = RGBColor(0x33, 0x33, 0x33)

def set_cn_font(run, name='微软雅黑', size=Pt(10.5)):
    run.font.name = name
    run.font.size = size
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '3A3F47')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    # left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '3A3F47')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_cn_font(run, '微软雅黑', Pt(14))
    run.font.bold = True
    run.font.color.rgb = DARK

def add_skill_line(doc, label, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(label)
    set_cn_font(run_label, '微软雅黑', Pt(10.5))
    run_label.font.bold = True
    run_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_items = p.add_run('    ' + items)
    set_cn_font(run_items, '微软雅黑', Pt(10.5))
    run_items.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('•  ' + text)
    set_cn_font(run, '微软雅黑', Pt(10.5))
    run.font.color.rgb = TEXT

def add_normal(doc, text, bold=False, color=TEXT, size=Pt(10.5)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_cn_font(run, '微软雅黑', size)
    run.font.bold = bold
    run.font.color.rgb = color

# ============================================================
# 内容
# ============================================================

# --- 头部 ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('——')
set_cn_font(run, '微软雅黑', Pt(24))
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run('Java后端开发 / 全栈开发工程师')
set_cn_font(run, '微软雅黑', Pt(11))
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('手机：——  |  邮箱：——  |  出生年月：——  |  现居：——')
set_cn_font(run, '微软雅黑', Pt(10.5))
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_line(doc)

# --- 教育背景 ---
add_section_title(doc, '教育背景')
add_normal(doc, '2022.09 - 2026.06     ——大学    计算机科学与技术    本科')

# --- 技术栈 ---
add_section_title(doc, '技术栈')
add_skill_line(doc, '后端开发：', 'Java，Spring Boot 3，MyBatis，MySQL，RESTful API 设计，Maven')
add_skill_line(doc, '前端开发：', 'Vue 3（组合式API），Element Plus，JavaScript / ES6，Axios，Vite')
add_skill_line(doc, '数据可视化：', 'ECharts（力导向图、柱状图、饼图等图表类型）')
add_skill_line(doc, '安全认证：', 'JWT 无状态认证，BCrypt 密码加密，Spring MVC 拦截器鉴权')
add_skill_line(doc, 'AI 集成：', '大语言模型 API 调用，提示词工程，流式响应处理')
add_skill_line(doc, '开发工具：', 'IntelliJ IDEA，VS Code，Git，Postman')

# --- 项目经历 ---
add_section_title(doc, '项目经历')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('智能笔记系统（个人博客 + 知识管理平台）')
set_cn_font(run, '微软雅黑', Pt(12))
run.font.bold = True
run.font.color.rgb = DARK
run2 = p.add_run('    2025.09 - 2026.05 | 独立全栈开发')
set_cn_font(run2, '微软雅黑', Pt(9))
run2.font.color.rgb = GRAY

add_normal(doc, '设计并开发了一个基于 SpringBoot + Vue3 的个人智能笔记系统，实现从内容创作到知识沉淀再到数据分析的全流程闭环。系统集成大语言模型提供 AI 辅助写作，通过知识图谱与写作仪表盘实现个人知识资产的结构化分析与可视化呈现。')

add_normal(doc, '技术架构：SpringBoot 3 + MyBatis + MySQL（后端）| Vue 3 + Element Plus + ECharts（前端）| JWT + BCrypt（安全）', color=RGBColor(0x88, 0x88, 0x88), size=Pt(9))

bullets = [
    '独立完成系统前后端架构设计与编码，采用 RESTful 风格设计 50+ 个接口，实现前后端分离开发与联调',
    '基于 MyBatis 实现复杂查询，包括分类条件筛选、LIKE 模糊匹配搜索、手动 LIMIT 分页及多表 JOIN 聚合统计',
    '实现 JWT 无状态令牌认证机制，通过拦截器实现接口级权限控制，登录注册密码经 BCrypt 加密存储',
    '集成大语言模型 API，通过统一接口按动作类型分发，实现文章续写、内容总结、标题优化与错误检测四项 AI 辅助功能',
    '采用 Markdown 源文本存储，前端实时渲染预览，后端支持 MD / TXT / EPUB 三种格式的文件流式生成与下载',
    '基于标题滑动窗口分词与停用词过滤构建关键词关联网络，使用 ECharts 力导向图实现知识图谱可视化',
    '通过 SQL 聚合查询实现写作仪表盘，以柱状图和饼图展示发文趋势与分类占比等多维度统计数据',
    '使用 Vue3 组合式 API 与 Pinia 管理全局状态，配合 localStorage 实现 Token 持久化与草稿暂存',
    '编写 60+ 组功能测试用例，覆盖正常流程与异常边界，各模块均通过验证',
]
for b in bullets:
    add_bullet(doc, b)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run('项目成果：')
set_cn_font(run, '微软雅黑', Pt(10.5))
run.font.bold = True
run.font.color.rgb = DARK
run2 = p.add_run('系统涵盖用户认证、文章浏览搜索、评论互动、点赞收藏、多格式下载、AI 辅助写作、知识图谱可视化及写作仪表盘六大核心模块，实现了数据自主可控的完整知识管理闭环。')
set_cn_font(run2, '微软雅黑', Pt(10.5))
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# --- 自我评价 ---
add_section_title(doc, '自我评价')
evals = [
    '具备独立完成从前端界面到后端服务到数据库设计的全栈开发能力。',
    '熟悉 SpringBoot + Vue3 主流技术栈，有完整的项目开发、测试与部署经验。',
    '善于通过技术手段解决实际问题，对 AI 技术应用有兴趣和实践经验。',
    '具备良好的代码规范意识与文档编写习惯，能够独立完成毕业论文撰写与答辩准备。',
    '学习能力强，能快速上手新技术并应用到实际项目中。',
]
for e in evals:
    add_bullet(doc, e)

# ===== 保存 =====
output_path = os.path.join(os.path.dirname(__file__), '个人简历.docx')
doc.save(output_path)
print(f'Word saved to: {output_path}')
